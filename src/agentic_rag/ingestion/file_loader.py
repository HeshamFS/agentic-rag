"""
File loader for document ingestion.

Supports multiple file formats: PDF, DOCX, TXT, MD, HTML.
"""

from pathlib import Path

from pydantic import BaseModel

from agentic_rag.core.models import Document


class LoadResult(BaseModel):
    """Result of loading a file."""

    document: Document | None = None
    success: bool = True
    error: str | None = None
    file_path: str = ""
    file_type: str = ""


class FileLoader:
    """
    Multi-format file loader.

    Supports:
    - Plain text (.txt)
    - Markdown (.md)
    - PDF (.pdf) - requires pypdf
    - Word (.docx) - requires python-docx
    - HTML (.html, .htm) - requires beautifulsoup4
    """

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx", ".html", ".htm"}

    def __init__(self, encoding: str = "utf-8"):
        """
        Initialize file loader.

        Args:
            encoding: Default text encoding.
        """
        self._encoding = encoding

    def load(self, file_path: str | Path) -> LoadResult:
        """
        Load a single file.

        Args:
            file_path: Path to file.

        Returns:
            LoadResult with document or error.
        """
        path = Path(file_path)

        if not path.exists():
            return LoadResult(
                success=False,
                error=f"File not found: {path}",
                file_path=str(path),
            )

        if not path.is_file():
            return LoadResult(
                success=False,
                error=f"Not a file: {path}",
                file_path=str(path),
            )

        ext = path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            return LoadResult(
                success=False,
                error=f"Unsupported file type: {ext}",
                file_path=str(path),
                file_type=ext,
            )

        try:
            if ext in {".txt", ".md"}:
                content = self._load_text(path)
            elif ext == ".pdf":
                content = self._load_pdf(path)
            elif ext == ".docx":
                content = self._load_docx(path)
            elif ext in {".html", ".htm"}:
                content = self._load_html(path)
            else:
                content = self._load_text(path)

            document = Document(
                content=content,
                source=str(path),
                metadata={
                    "source": str(path),
                    "filename": path.name,
                    "extension": ext,
                    "size_bytes": path.stat().st_size,
                },
            )

            return LoadResult(
                document=document,
                success=True,
                file_path=str(path),
                file_type=ext,
            )

        except Exception as e:
            return LoadResult(
                success=False,
                error=str(e),
                file_path=str(path),
                file_type=ext,
            )

    def _load_text(self, path: Path) -> str:
        """Load plain text file."""
        return path.read_text(encoding=self._encoding)

    def _load_pdf(self, path: Path) -> str:
        """Load PDF file."""
        try:
            from pypdf import PdfReader
        except ImportError:
            raise ImportError("pypdf required for PDF support: pip install pypdf")

        reader = PdfReader(str(path))
        text_parts = []

        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)

        return "\n\n".join(text_parts)

    def _load_docx(self, path: Path) -> str:
        """Load Word document."""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx required for DOCX support: pip install python-docx")

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n\n".join(paragraphs)

    def _load_html(self, path: Path) -> str:
        """Load HTML file and extract text."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError(
                "beautifulsoup4 required for HTML support: pip install beautifulsoup4"
            )

        html = path.read_text(encoding=self._encoding)
        soup = BeautifulSoup(html, "html.parser")

        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()

        # Get text
        text = soup.get_text(separator="\n")

        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = "\n".join(chunk for chunk in chunks if chunk)

        return text

    def load_directory(
        self,
        directory: str | Path,
        recursive: bool = True,
        extensions: set[str] | None = None,
    ) -> list[LoadResult]:
        """
        Load all supported files from a directory.

        Args:
            directory: Directory path.
            recursive: Search subdirectories.
            extensions: Filter by extensions (default: all supported).

        Returns:
            List of LoadResult for each file.
        """
        path = Path(directory)
        if not path.is_dir():
            return [
                LoadResult(
                    success=False,
                    error=f"Not a directory: {path}",
                    file_path=str(path),
                )
            ]

        extensions = extensions or self.SUPPORTED_EXTENSIONS
        results = []

        pattern = "**/*" if recursive else "*"
        for file_path in path.glob(pattern):
            if file_path.is_file() and file_path.suffix.lower() in extensions:
                result = self.load(file_path)
                results.append(result)

        return results


class URLLoader:
    """
    Load documents from URLs.

    Fetches web pages and extracts text content.
    """

    def __init__(self, timeout: float = 30.0):
        """
        Initialize URL loader.

        Args:
            timeout: Request timeout in seconds.
        """
        self._timeout = timeout

    async def load(self, url: str) -> LoadResult:
        """
        Load document from URL.

        Args:
            url: URL to fetch.

        Returns:
            LoadResult with document or error.
        """
        try:
            import httpx
            from bs4 import BeautifulSoup
        except ImportError:
            raise ImportError("httpx and beautifulsoup4 required: pip install httpx beautifulsoup4")

        try:
            async with httpx.AsyncClient() as client:
                response = await client.get(url, timeout=self._timeout, follow_redirects=True)
                response.raise_for_status()

            # Parse HTML
            soup = BeautifulSoup(response.text, "html.parser")

            # Remove unwanted elements
            for tag in soup(["script", "style", "nav", "footer", "header"]):
                tag.decompose()

            # Extract title
            title = soup.title.string if soup.title else url

            # Extract main content
            main = soup.find("main") or soup.find("article") or soup.body
            text = main.get_text(separator="\n") if main else soup.get_text(separator="\n")

            # Clean whitespace
            lines = (line.strip() for line in text.splitlines())
            text = "\n".join(line for line in lines if line)

            document = Document(
                content=text,
                source=url,
                metadata={
                    "source": url,
                    "title": title,
                    "content_type": response.headers.get("content-type", ""),
                },
            )

            return LoadResult(
                document=document,
                success=True,
                file_path=url,
                file_type="url",
            )

        except Exception as e:
            return LoadResult(
                success=False,
                error=str(e),
                file_path=url,
                file_type="url",
            )
